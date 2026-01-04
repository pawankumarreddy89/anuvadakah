from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.concurrency import run_in_threadpool 
from pydantic import BaseModel
from transformers import pipeline 
import torch
import shutil
import os
import uuid
import re 
import gc 
import platform # 🟢 NEW: To detect Windows vs Linux
from PIL import Image
import pytesseract
import pdfplumber
from docx import Document
from gtts import gTTS 

# PDF LIBRARIES
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

app = FastAPI()

# 🟢 DYNAMIC TESSERACT CONFIGURATION
if platform.system() == "Windows":
    # YOUR LOCAL PATH
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
else:
    # LINUX CLOUD PATH (Standard)
    pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

print("------------------------------------------------")
print(f"🚀 STARTING ANUVADAKAH ON {platform.system().upper()}...")
print("------------------------------------------------")

# LOAD AI
try:
    device = 0 if torch.cuda.is_available() else -1 
    translator_pipe = pipeline("translation", model="facebook/nllb-200-distilled-600M", device=device)
    print(f"✅ AI ENGINE ONLINE!")
except Exception as e:
    print(f"❌ AI FAILURE: {e}")

LANG_MAP = {
    "hi": "hin_Deva", "te": "tel_Telu", "ta": "tam_Taml", "kn": "kan_Knda",
    "ml": "mal_Mlym", "bn": "ben_Beng", "mr": "mar_Deva", "gu": "guj_Gujr",
    "pa": "pan_Guru", "or": "ory_Orya", "as": "asm_Beng", "ur": "urd_Arab",
    "sa": "san_Deva", "en": "eng_Latn", "fr": "fra_Latn"
}

# ==========================================
# 🟢 HYBRID FONT LOADER (Windows & Linux)
# ==========================================
FONT_NAME = 'Helvetica' # Default

def register_fonts():
    global FONT_NAME
    font_candidates = []

    if platform.system() == "Windows":
        # Windows Fonts
        font_candidates = [
            ("Nirmala", r"C:\Windows\Fonts\Nirmala.ttc"),
            ("Nirmala", r"C:\Windows\Fonts\Nirmala.ttf"),
            ("ArialUnicode", r"C:\Windows\Fonts\arialuni.ttf")
        ]
    else:
        # Linux/Cloud Fonts (We will install these in Docker)
        font_candidates = [
            ("NotoSans", "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf"),
            ("NotoSans", "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf"),
            ("FreeSerif", "/usr/share/fonts/truetype/freefont/FreeSerif.ttf")
        ]

    for name, path in font_candidates:
        if os.path.exists(path):
            try:
                if path.endswith(".ttc"):
                    pdfmetrics.registerFont(TTFont(name, path)) # Try simple load for TTC
                else:
                    pdfmetrics.registerFont(TTFont(name, path))
                FONT_NAME = name
                print(f"✅ Loaded Font: {name} from {path}")
                return
            except Exception as e:
                print(f"⚠️ Failed to load {path}: {e}")
    
    print("⚠️ Using Default Helvetica (No Hindi Support)")

register_fonts()

def clean_for_xml(text):
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)

def split_text_safe(text, max_chars=400):
    chunks = []
    while len(text) > max_chars:
        split_at = text.rfind(' ', 0, max_chars)
        if split_at == -1: split_at = max_chars
        chunks.append(text[:split_at])
        text = text[split_at:]
    chunks.append(text)
    return chunks

def process_translation_heavy(text, target_code):
    output = translator_pipe(text, src_lang="eng_Latn", tgt_lang=target_code)
    return output[0]['translation_text']

def process_pdf_heavy(path):
    full_text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text: full_text += text + "\n"
    return full_text

def generate_pdf_file(text):
    filename = f"translation_{uuid.uuid4()}.pdf"
    c = canvas.Canvas(filename, pagesize=A4)
    width, height = A4
    c.setFont(FONT_NAME, 16)
    c.drawString(50, height - 50, "Anuvadakah Translation")
    c.line(50, height - 60, width - 50, height - 60)
    text_object = c.beginText(50, height - 100)
    text_object.setFont(FONT_NAME, 12)
    lines = text.split('\n')
    for line in lines:
        try:
            wrapped_lines = simpleSplit(line, FONT_NAME, 12, width - 100)
        except:
            wrapped_lines = [line] 
        for wrapped in wrapped_lines:
            text_object.textLine(wrapped)
            if text_object.getY() < 50:
                c.drawText(text_object)
                c.showPage()
                text_object = c.beginText(50, height - 50)
                text_object.setFont(FONT_NAME, 12)
    c.drawText(text_object)
    c.save()
    return filename

class TranslationRequest(BaseModel):
    text: str
    target_lang: str

class SpeakRequest(BaseModel):
    text: str
    lang: str

@app.get("/")
def read_root(): return {"message": "Anuvādakaḥ Online"}

@app.post("/translate")
async def translate_text(request: TranslationRequest):
    try:
        if not request.text.strip(): return {"translated_text": "", "status": "success"}
        target_code = LANG_MAP.get(request.target_lang, "eng_Latn")
        translated_text = await run_in_threadpool(process_translation_heavy, request.text, target_code)
        return {"translated_text": translated_text, "status": "success"}
    except Exception as e:
        return {"translated_text": f"Error: {str(e)}", "status": "error"}

@app.post("/pdf-translate")
async def pdf_translate(target_lang: str = Form(...), file: UploadFile = File(...)):
    temp_pdf = f"temp_{uuid.uuid4()}.pdf"
    try:
        with open(temp_pdf, "wb") as buffer: shutil.copyfileobj(file.file, buffer)
        extracted_text = await run_in_threadpool(process_pdf_heavy, temp_pdf)
        os.remove(temp_pdf)
        if not extracted_text.strip(): return {"original_text": "Scanned PDF", "translated_text": "Use Image Scan.", "status": "error"}
        target_code = LANG_MAP.get(target_lang, "eng_Latn")
        chunks = split_text_safe(extracted_text)
        translated_chunks = []
        for i, chunk in enumerate(chunks):
            t_text = await run_in_threadpool(process_translation_heavy, chunk, target_code)
            translated_chunks.append(t_text)
            if i % 10 == 0: gc.collect()
        full_translation = "\n".join(translated_chunks)
        gc.collect() 
        return {"original_text": "Processed", "translated_text": full_translation, "status": "success"}
    except Exception as e:
        if os.path.exists(temp_pdf): os.remove(temp_pdf)
        return {"status": "error", "message": str(e)}

@app.post("/image-translate")
async def image_translate(target_lang: str = Form(...), file: UploadFile = File(...)):
    try:
        image = Image.open(file.file)
        extracted_text = await run_in_threadpool(pytesseract.image_to_string, image)
        target_code = LANG_MAP.get(target_lang, "eng_Latn")
        translated_text = await run_in_threadpool(process_translation_heavy, extracted_text, target_code)
        return {"original_text": extracted_text, "translated_text": translated_text, "status": "success"}
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/download-docx")
async def download_docx(request: SpeakRequest):
    try:
        def make_doc():
            doc = Document()
            doc.add_heading('Anuvādakaḥ Translation', 0)
            doc.add_paragraph(clean_for_xml(request.text))
            fname = f"translation_{uuid.uuid4()}.docx"
            doc.save(fname)
            return fname
        filename = await run_in_threadpool(make_doc)
        return FileResponse(filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename)
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/download-pdf")
async def download_pdf(request: SpeakRequest):
    try:
        filename = await run_in_threadpool(generate_pdf_file, request.text)
        return FileResponse(filename, media_type="application/pdf", filename=filename)
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/speak")
async def speak_text(request: SpeakRequest):
    try:
        def make_audio():
            tts = gTTS(text=request.text, lang=request.lang)
            fname = f"audio_{uuid.uuid4()}.mp3"
            tts.save(fname)
            return fname
        filename = await run_in_threadpool(make_audio)
        return FileResponse(filename, media_type="audio/mpeg", filename="audio.mp3")
    except Exception as e:
        return {"status": "error", "message": str(e)}